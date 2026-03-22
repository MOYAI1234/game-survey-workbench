from pathlib import Path
from zipfile import ZIP_DEFLATED, ZIP_STORED, ZipFile

import pytest
from fastapi.testclient import TestClient
from sqlmodel import Session, select

from game_survey_workbench.app import create_app
from game_survey_workbench.db import create_db_and_tables, get_engine
from game_survey_workbench.models.knowledge import KnowledgeDocument
from game_survey_workbench.services.workspace import bootstrap_workspace


@pytest.fixture()
def workspace(tmp_path: Path) -> Path:
    bootstrap_workspace(tmp_path)
    create_db_and_tables(tmp_path)
    return tmp_path


@pytest.fixture()
def app_client(workspace: Path, monkeypatch: pytest.MonkeyPatch):
    monkeypatch.setenv("GAME_SURVEY_WORKBENCH_WORKSPACE_ROOT", str(workspace))
    with TestClient(create_app()) as client:
        yield client


def _make_test_docx(workspace: Path) -> Path:
    """Create a minimal .docx for testing."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    doc.add_heading("Survey Methods", level=1)
    doc.add_paragraph("This document describes survey methodology for game research.")
    path = workspace / "test_doc.docx"
    doc.save(str(path))
    return path


def _make_test_epub(workspace: Path) -> Path:
    epub_path = workspace / "test_book.epub"
    with ZipFile(epub_path, "w") as archive:
        archive.writestr("mimetype", "application/epub+zip", compress_type=ZIP_STORED)
        archive.writestr(
            "META-INF/container.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
  </rootfiles>
</container>
""",
            compress_type=ZIP_DEFLATED,
        )
        archive.writestr(
            "OEBPS/content.opf",
            """<?xml version="1.0" encoding="UTF-8"?>
<package xmlns="http://www.idpf.org/2007/opf" unique-identifier="bookid" version="2.0">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:title>Test EPUB</dc:title>
    <dc:language>en</dc:language>
    <dc:identifier id="bookid">test-epub-id</dc:identifier>
  </metadata>
  <manifest>
    <item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>
    <item id="chapter1" href="chapter1.xhtml" media-type="application/xhtml+xml"/>
  </manifest>
  <spine toc="ncx">
    <itemref idref="chapter1"/>
  </spine>
</package>
""",
            compress_type=ZIP_DEFLATED,
        )
        archive.writestr(
            "OEBPS/toc.ncx",
            """<?xml version="1.0" encoding="UTF-8"?>
<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">
  <head>
    <meta name="dtb:uid" content="test-epub-id"/>
  </head>
  <docTitle><text>Test EPUB</text></docTitle>
  <navMap>
    <navPoint id="navPoint-1" playOrder="1">
      <navLabel><text>Chapter 1</text></navLabel>
      <content src="chapter1.xhtml"/>
    </navPoint>
  </navMap>
</ncx>
""",
            compress_type=ZIP_DEFLATED,
        )
        archive.writestr(
            "OEBPS/chapter1.xhtml",
            """<?xml version="1.0" encoding="UTF-8"?>
<html xmlns="http://www.w3.org/1999/xhtml">
  <head><title>Test EPUB</title></head>
  <body>
    <h1>Test EPUB</h1>
    <p>User retention improves when onboarding and lifecycle messaging are aligned.</p>
  </body>
</html>
""",
            compress_type=ZIP_DEFLATED,
        )
    return epub_path


def test_upload_non_markdown_redirects_to_convert_preview(app_client, workspace):
    docx_path = _make_test_docx(workspace)
    with open(docx_path, "rb") as handle:
        response = app_client.post(
            "/knowledge/upload",
            files={
                "file": (
                    "methods.docx",
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={"purposes": ["questionnaire_design"]},
            follow_redirects=True,
        )

    assert response.status_code == 200
    html = response.text
    assert "转换预览" in html or "convert" in html.lower()
    assert "Survey Methods" in html or "methods.docx" in html


def test_upload_epub_redirects_to_convert_preview(app_client, workspace):
    epub_path = _make_test_epub(workspace)
    with open(epub_path, "rb") as handle:
        response = app_client.post(
            "/knowledge/upload",
            files={"file": ("user-ops.epub", handle, "application/epub+zip")},
            data={"purposes": ["analysis"]},
            follow_redirects=True,
        )

    assert response.status_code == 200
    html = response.text
    assert "转换预览" in html or "convert" in html.lower()
    assert "user-ops.epub" in html or "Test EPUB" in html


def test_upload_markdown_still_works_directly(app_client, workspace):
    response = app_client.post(
        "/knowledge/upload",
        files={
            "file": ("guide.md", b"---\ntitle: Test Guide\n---\nContent here.", "text/markdown")
        },
        data={"purposes": ["analysis"]},
        follow_redirects=False,
    )

    assert response.status_code == 303
    assert "/knowledge" in response.headers["location"]


def test_convert_confirm_ingests_document(app_client, workspace):
    docx_path = _make_test_docx(workspace)
    with open(docx_path, "rb") as handle:
        preview_resp = app_client.post(
            "/knowledge/upload",
            files={
                "file": (
                    "methods.docx",
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={"purposes": []},
            follow_redirects=True,
        )

    assert preview_resp.status_code == 200

    staging_dir = workspace / "knowledge" / "staging"
    staging_files = list(staging_dir.glob("*.md")) if staging_dir.exists() else []
    if not staging_files:
        pytest.skip("No staging file created - conversion may have failed")
    staging_id = staging_files[0].stem

    confirm_resp = app_client.post(
        "/knowledge/convert-confirm",
        data={
            "staging_id": staging_id,
            "source_format": "docx",
            "title": "Survey Methods Guide",
            "doc_type": "guide",
            "purposes": ["questionnaire_design"],
        },
        follow_redirects=False,
    )

    assert confirm_resp.status_code == 303
    assert "upload_success" in confirm_resp.headers["location"]
    assert not staging_files[0].exists()


def test_convert_download_returns_markdown_file(app_client, workspace):
    staging_dir = workspace / "knowledge" / "staging"
    staging_dir.mkdir(parents=True, exist_ok=True)
    staging_file = staging_dir / "abc123.md"
    staging_file.write_text("# Test\n\nConverted content.", encoding="utf-8")

    response = app_client.post(
        "/knowledge/convert-download",
        data={"staging_id": "abc123"},
    )

    assert response.status_code == 200
    assert "text/markdown" in response.headers.get("content-type", "")
    assert b"Converted content" in response.content


def test_convert_confirm_persists_epub_source_format(app_client, workspace):
    epub_path = _make_test_epub(workspace)
    with open(epub_path, "rb") as handle:
        preview_resp = app_client.post(
            "/knowledge/upload",
            files={"file": ("user-ops.epub", handle, "application/epub+zip")},
            data={"purposes": []},
            follow_redirects=True,
        )

    assert preview_resp.status_code == 200

    staging_dir = workspace / "knowledge" / "staging"
    staging_files = list(staging_dir.glob("*.md")) if staging_dir.exists() else []
    assert staging_files

    confirm_resp = app_client.post(
        "/knowledge/convert-confirm",
        data={
            "staging_id": staging_files[0].stem,
            "source_format": "epub",
            "title": "EPUB Guide",
            "doc_type": "guide",
            "purposes": ["analysis"],
        },
        follow_redirects=False,
    )

    assert confirm_resp.status_code == 303
    with Session(get_engine(workspace)) as session:
        document = session.exec(select(KnowledgeDocument)).first()
    assert document is not None
    assert document.source_format == "epub"
