from pathlib import Path

import pytest

from game_survey_workbench.services.knowledge_convert import (
    SUPPORTED_CONVERSION_EXTENSIONS,
    ConversionResult,
    assess_conversion_quality,
    convert_to_markdown,
)


def test_supported_extensions_includes_pdf_docx_pptx():
    assert ".pdf" in SUPPORTED_CONVERSION_EXTENSIONS
    assert ".docx" in SUPPORTED_CONVERSION_EXTENSIONS
    assert ".pptx" in SUPPORTED_CONVERSION_EXTENSIONS
    assert ".md" not in SUPPORTED_CONVERSION_EXTENSIONS


def test_convert_to_markdown_converts_docx(tmp_path: Path):
    """Create a minimal .docx via python-docx and convert it."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    doc.add_heading("Test Title", level=1)
    doc.add_paragraph("This is a test paragraph about game survey methodology.")
    docx_path = tmp_path / "test.docx"
    doc.save(str(docx_path))

    result = convert_to_markdown(docx_path)

    assert isinstance(result, ConversionResult)
    assert result.success is True
    assert "Test Title" in result.markdown_text
    assert "test paragraph" in result.markdown_text
    assert result.error_message is None


def test_convert_to_markdown_returns_failure_for_unsupported_format(tmp_path: Path):
    bad_file = tmp_path / "data.zip"
    bad_file.write_bytes(b"PK\x03\x04fake zip content")

    result = convert_to_markdown(bad_file)

    assert result.success is False
    assert result.error_message is not None


def test_convert_to_markdown_handles_empty_file(tmp_path: Path):
    empty = tmp_path / "empty.pdf"
    empty.write_bytes(b"")

    result = convert_to_markdown(empty)

    if result.success:
        assert result.markdown_text is not None


def test_assess_conversion_quality_flags_empty_text():
    quality = assess_conversion_quality("")
    assert quality.warning is not None
    assert "空" in quality.warning or "empty" in quality.warning.lower()


def test_assess_conversion_quality_flags_garbled_text():
    garbled = "正常文字" + "\ufffd" * 20 + "abc"
    quality = assess_conversion_quality(garbled)
    assert quality.is_low_quality is True


def test_assess_conversion_quality_passes_clean_text():
    clean = "This is a clean document about game survey design methodology. " * 10
    quality = assess_conversion_quality(clean)
    assert quality.is_low_quality is False
    assert quality.warning is None
